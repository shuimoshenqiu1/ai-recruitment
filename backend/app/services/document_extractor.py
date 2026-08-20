"""文档文本提取服务 - 支持PDF/DOCX/DOC/TXT格式"""

import logging
import os
import shutil
import subprocess
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

from app.core.config import settings

logger = logging.getLogger(__name__)

# === 安全限制常量 ===
MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024  # 50MB
MAX_PDF_PAGES = 200
MAX_TEXT_LENGTH = 500 * 1024  # 500KB文本


class ExtractionError(Exception):
    """文本提取失败"""

    def __init__(self, message: str, file_type: str = "unknown"):
        self.message = message
        self.file_type = file_type
        super().__init__(message)


def _validate_file_path(file_path: str) -> None:
    """验证文件路径是否在允许的目录内，防止路径遍历攻击。

    Args:
        file_path: 待验证的文件路径

    Raises:
        ExtractionError: 路径不在允许范围内
    """
    allowed_base = os.path.realpath(settings.UPLOAD_DIR)
    real_path = os.path.realpath(file_path)

    if not real_path.startswith(allowed_base + os.sep) and real_path != allowed_base:
        raise ExtractionError(
            f"文件路径不在允许范围内: {file_path}",
            file_type="security",
        )


def _check_file_size(file_path: str, file_extension: str) -> None:
    """检查文件大小是否超过限制。

    Args:
        file_path: 文件路径
        file_extension: 文件扩展名

    Raises:
        ExtractionError: 文件超过大小限制
    """
    file_size = os.path.getsize(file_path)
    if file_size > MAX_FILE_SIZE_BYTES:
        raise ExtractionError(
            f"文件过大({file_size // 1024 // 1024}MB)，超过限制(50MB)",
            file_type=file_extension,
        )


# ============================================================
# PDF 提取
# ============================================================


async def extract_from_pdf(file_path: str) -> str:
    """
    从PDF提取文本，保留基本结构。

    使用PyMuPDF（fitz）逐页提取纯文本，页间以双换行分隔。

    Args:
        file_path: PDF文件路径

    Returns:
        提取的纯文本

    Raises:
        ExtractionError: 文件无法打开或解析失败
    """
    if not Path(file_path).exists():
        raise ExtractionError(f"文件不存在: {file_path}", file_type="pdf")

    try:
        doc = fitz.open(file_path)
    except Exception as e:
        raise ExtractionError(f"无法打开PDF文件: {e}", file_type="pdf")

    try:
        text_parts: list[str] = []
        total_pages = len(doc)
        pages_to_process = min(total_pages, MAX_PDF_PAGES)

        if total_pages > MAX_PDF_PAGES:
            logger.warning(
                f"PDF页数({total_pages})超过限制({MAX_PDF_PAGES})，仅处理前{MAX_PDF_PAGES}页"
            )

        total_text_len = 0
        for page_num, page in enumerate(doc, start=1):
            if page_num > pages_to_process:
                break
            try:
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(text)
                    total_text_len += len(text)
                    if total_text_len > MAX_TEXT_LENGTH:
                        logger.warning(
                            f"PDF文本量({total_text_len}字符)超过限制({MAX_TEXT_LENGTH})，截断"
                        )
                        break
            except Exception as e:
                logger.warning(f"PDF第{page_num}页提取失败: {e}")
                continue

        result_text = "\n\n".join(text_parts).strip()
        # 最终截断保护
        if len(result_text) > MAX_TEXT_LENGTH:
            result_text = result_text[:MAX_TEXT_LENGTH]
        return result_text
    finally:
        doc.close()


# ============================================================
# DOCX 提取
# ============================================================


async def extract_from_docx(file_path: str) -> str:
    """
    从DOCX提取文本，保留段落结构并包含表格内容。

    Args:
        file_path: DOCX文件路径

    Returns:
        提取的纯文本（段落+表格）

    Raises:
        ExtractionError: 文件无法打开或解析失败
    """
    if not Path(file_path).exists():
        raise ExtractionError(f"文件不存在: {file_path}", file_type="docx")

    try:
        doc = Document(file_path)
    except Exception as e:
        raise ExtractionError(f"无法打开DOCX文件: {e}", file_type="docx")

    paragraphs: list[str] = []

    # 提取正文段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                # 去重相邻相同单元格（合并单元格会导致重复）
                deduped = []
                for cell_text in cells:
                    if not deduped or deduped[-1] != cell_text:
                        deduped.append(cell_text)
                paragraphs.append(" | ".join(deduped))

    return "\n".join(paragraphs).strip()


# ============================================================
# DOC 提取（旧版Word格式）
# ============================================================


async def extract_from_doc(file_path: str) -> str:
    """
    从DOC（旧版Word）提取文本。

    优先使用antiword命令行工具，如果不可用则尝试LibreOffice转换。
    两者都不可用时返回明确错误信息。

    Args:
        file_path: DOC文件路径

    Returns:
        提取的纯文本

    Raises:
        ExtractionError: 无法提取或工具不可用
    """
    if not Path(file_path).exists():
        raise ExtractionError(f"文件不存在: {file_path}", file_type="doc")

    # 方案1：尝试 antiword
    text = _try_antiword(file_path)
    if text is not None:
        return text

    # 方案2：尝试 LibreOffice 转换为纯文本
    text = _try_libreoffice(file_path)
    if text is not None:
        return text

    raise ExtractionError(
        "不支持.doc格式提取：系统未安装antiword或LibreOffice。"
        "建议将文件转换为.docx格式后重新上传。",
        file_type="doc",
    )


def _try_antiword(file_path: str) -> str | None:
    """尝试使用antiword提取DOC文本"""
    if not shutil.which("antiword"):
        logger.debug("antiword未安装，跳过")
        return None

    try:
        result = subprocess.run(
            ["antiword", "-w", "0", file_path],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
        logger.warning(f"antiword返回码: {result.returncode}, stderr: {result.stderr}")
        return None
    except (subprocess.TimeoutExpired, OSError) as e:
        logger.warning(f"antiword执行失败: {e}")
        return None


def _try_libreoffice(file_path: str) -> str | None:
    """尝试使用LibreOffice将DOC转换为文本"""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        logger.debug("LibreOffice未安装，跳过")
        return None

    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            result = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    tmp_dir,
                    file_path,
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )

            if result.returncode != 0:
                logger.warning(f"LibreOffice转换失败: {result.stderr}")
                return None

            # 查找生成的txt文件
            txt_files = list(Path(tmp_dir).glob("*.txt"))
            if not txt_files:
                logger.warning("LibreOffice未生成txt文件")
                return None

            text = txt_files[0].read_text(encoding="utf-8", errors="replace")
            return text.strip() if text.strip() else None
    except (subprocess.TimeoutExpired, OSError) as e:
        logger.warning(f"LibreOffice执行失败: {e}")
        return None


# ============================================================
# TXT 提取（多编码自动检测）
# ============================================================

_TXT_ENCODINGS = ["utf-8", "gbk", "gb2312", "gb18030", "big5", "latin-1"]


async def extract_from_txt(file_path: str) -> str:
    """
    从TXT文件提取文本，自动检测编码。

    按优先级尝试：UTF-8 -> GBK -> GB2312 -> GB18030 -> Big5 -> Latin-1
    Latin-1作为兜底（永远不会失败）。

    Args:
        file_path: TXT文件路径

    Returns:
        提取的纯文本

    Raises:
        ExtractionError: 文件不存在或读取失败
    """
    path = Path(file_path)
    if not path.exists():
        raise ExtractionError(f"文件不存在: {file_path}", file_type="txt")

    raw_bytes = path.read_bytes()

    # 检查BOM
    if raw_bytes.startswith(b"\xef\xbb\xbf"):
        return raw_bytes[3:].decode("utf-8", errors="replace").strip()
    if raw_bytes.startswith(b"\xff\xfe"):
        return raw_bytes[2:].decode("utf-16-le", errors="replace").strip()
    if raw_bytes.startswith(b"\xfe\xff"):
        return raw_bytes[2:].decode("utf-16-be", errors="replace").strip()

    # 按优先级尝试编码
    for encoding in _TXT_ENCODINGS:
        try:
            text = raw_bytes.decode(encoding)
            # 简单校验：解码后不应包含大量替换字符
            if "\ufffd" not in text:
                return text.strip()
        except (UnicodeDecodeError, LookupError):
            continue

    # 最终兜底：latin-1永远成功
    return raw_bytes.decode("latin-1").strip()


# ============================================================
# 图片OCR预留
# ============================================================


async def extract_from_image(file_path: str) -> str:
    """
    从图片提取文本（OCR）- 功能开发中。

    当前为占位实现，后续将集成OCR引擎（如Tesseract/PaddleOCR）。

    Args:
        file_path: 图片文件路径

    Returns:
        占位信息文本
    """
    if not Path(file_path).exists():
        raise ExtractionError(f"文件不存在: {file_path}", file_type="image")

    logger.info(f"图片OCR功能开发中，无法提取: {file_path}")
    return "[OCR功能开发中] 当前版本暂不支持从图片中提取文本，请上传PDF或DOCX格式的简历。"


# ============================================================
# 统一入口
# ============================================================

# 支持的提取器映射
_EXTRACTORS: dict[str, callable] = {
    "pdf": extract_from_pdf,
    "docx": extract_from_docx,
    "doc": extract_from_doc,
    "txt": extract_from_txt,
    # 图片格式（实验性）
    "jpg": extract_from_image,
    "jpeg": extract_from_image,
    "png": extract_from_image,
}


async def extract_text(file_path: str, file_extension: str) -> tuple[str, dict]:
    """
    统一文本提取入口。

    根据文件扩展名路由到对应的提取器，返回提取的文本和元数据。

    Args:
        file_path: 文件路径
        file_extension: 文件扩展名（小写，不含点号）

    Returns:
        (text, metadata) 元组：
        - text: 提取的全文
        - metadata: 包含页数、字数、提取器等信息的字典

    Raises:
        ExtractionError: 不支持的格式或提取失败
    """
    ext = file_extension.lower().lstrip(".")

    # 安全检查：路径遍历防护 (C-1)
    _validate_file_path(file_path)

    # 安全检查：文件大小限制 (H-5)
    _check_file_size(file_path, ext)

    if ext not in _EXTRACTORS:
        supported = ", ".join(sorted(_EXTRACTORS.keys()))
        raise ExtractionError(
            f"不支持的文件格式: .{ext}，支持格式: {supported}",
            file_type=ext,
        )

    extractor = _EXTRACTORS[ext]
    logger.info(f"开始提取文本: file={file_path}, type={ext}")

    text = await extractor(file_path)

    # 构建元数据
    metadata = _build_metadata(text, ext, file_path)

    logger.info(
        f"文本提取完成: file={file_path}, "
        f"chars={metadata['char_count']}, lines={metadata['line_count']}"
    )

    return text, metadata


def _build_metadata(text: str, file_type: str, file_path: str) -> dict:
    """构建提取结果的元数据"""
    file_size = os.path.getsize(file_path) if Path(file_path).exists() else 0

    metadata = {
        "file_type": file_type,
        "file_size_bytes": file_size,
        "char_count": len(text),
        "line_count": text.count("\n") + 1 if text else 0,
        "extractor": f"extract_from_{file_type}",
        "is_ocr": file_type in ("jpg", "jpeg", "png"),
    }

    # PDF特有：页数
    if file_type == "pdf":
        try:
            doc = fitz.open(file_path)
            metadata["page_count"] = len(doc)
            doc.close()
        except Exception:
            metadata["page_count"] = None

    return metadata
